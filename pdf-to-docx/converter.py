import sys
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import traceback

def convert_pdf_to_docx(pdf_path, docx_path):
    try:
        print(f"[INFO] Starting PDF to DOCX conversion: {pdf_path}")
        
        if not os.path.exists(pdf_path):
            print(f"[ERROR] Input file not found: {pdf_path}")
            return False
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        
        print("[INFO] Opening PDF document...")
        pdf_document = fitz.open(pdf_path)
        doc = Document()
        
        total_pages = len(pdf_document)
        print(f"[INFO] Processing {total_pages} pages...")
        
        # Convert each page
        for page_num in range(total_pages):
            page = pdf_document.load_page(page_num)
            
            # Extract text from the page
            text = page.get_text()
            
            if text.strip():
                # Add a paragraph for each page
                p = doc.add_paragraph()
                p.add_run(f"--- Page {page_num + 1} ---").bold = True
                doc.add_paragraph(text)
            
            progress = int(((page_num + 1) / total_pages) * 100)
            print(f"[PROGRESS] Processed page {page_num + 1}/{total_pages} ({progress}%)")
        
        print("[INFO] Saving DOCX document...")
        doc.save(docx_path)
        pdf_document.close()
        
        # Verify the output file
        if os.path.exists(docx_path):
            file_size = os.path.getsize(docx_path)
            print(f"[SUCCESS] Conversion completed successfully: {docx_path} ({file_size} bytes)")
            return True
        else:
            print("[ERROR] Output file was not created")
            return False
            
    except Exception as e:
        print(f"[ERROR] Conversion failed: {str(e)}")
        traceback.print_exc()
        return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python converter.py <input_pdf> <output_docx>")
        sys.exit(1)
        
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    
    success = convert_pdf_to_docx(pdf_path, docx_path)
    sys.exit(0 if success else 1)