import sys
import os
import json
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTFigure
from docx import Document
from docx.shared import Inches
import io
from PIL import Image

def pdf_to_docx(pdf_path, output_dir):
    try:
        os.makedirs(output_dir, exist_ok=True)
        pdf_name = os.path.basename(pdf_path)
        docx_name = pdf_name.replace('.pdf', '.docx')
        docx_path = os.path.join(output_dir, docx_name)

        doc = Document()
        
        for page_layout in extract_pages(pdf_path):
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    # Extract text
                    text = element.get_text()
                    # Add paragraph to Word document
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(text)
                elif isinstance(element, LTFigure):
                    # Handle images (basic implementation)
                    try:
                        for img in element:
                            if hasattr(img, 'stream'):
                                img_data = img.stream.get_data()
                                image = Image.open(io.BytesIO(img_data))
                                img_path = os.path.join(output_dir, 'temp_img.png')
                                image.save(img_path)
                                doc.add_picture(img_path, width=Inches(4))
                                os.remove(img_path)
                    except Exception as img_error:
                        print(f"Image processing error: {img_error}")

        doc.save(docx_path)

        print(json.dumps({
            "status": "success",
            "output_path": docx_path
        }))
        return True
    except Exception as e:
        print(json.dumps({
            "status": "error",
            "message": str(e)
        }))
        return False

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(json.dumps({
            "status": "error",
            "message": "Usage: python script.py <pdf_path> <output_dir>"
        }))
        sys.exit(1)

    pdf_path = sys.argv[1]
    output_dir = sys.argv[2]

    if not os.path.exists(pdf_path):
        print(json.dumps({
            "status": "error",
            "message": f"File not found: {pdf_path}"
        }))
        sys.exit(1)

    success = pdf_to_docx(pdf_path, output_dir)
    sys.exit(0 if success else 1)